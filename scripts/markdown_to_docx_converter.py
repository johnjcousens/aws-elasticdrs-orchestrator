#!/usr/bin/env python3
"""
AWS Professional Markdown to DOCX Converter
===========================================

Converts Markdown files to professional AWS Word documents with proper formatting,
styles, and document structure suitable for AWS Professional Services documentation.

FEATURES:
---------
- **PROFESSIONAL AWS DOCUMENT STRUCTURE**: Automatic 6-page front matter generation
  • Title page with document title, subtitle, and date
  • Notices page with AWS disclaimers and copyright
  • Table of Contents pages (2 pages reserved for Word auto-generation)
  • Abstract page with document overview
  • Introduction page with strategic context
  • Content section starting with Executive Summary

- **AWS BRANDING AND FORMATTING**: Complete AWS Professional Services styling
  • Amazon Ember font throughout document
  • Amazon orange headings (RGB: 255, 153, 0)
  • Official AWS "Powered by AWS" logo integration (title page and footers)
  • Professional page headers starting on page 8
  • AWS Trademark Guidelines compliant branding

- **INTELLIGENT DOCUMENT DETECTION**: Automatic professional vs. generic conversion
  • Professional documents: Detected by presence of Executive Summary + title structure
  • Generic documents: Clean markdown conversion without front matter
  • Smart section processing with front matter integration

- **ADVANCED MARKDOWN PROCESSING**: Complete markdown element support
  • Headers (# ## ### #### #####) with proper hierarchy
  • Code blocks (```code```) with Consolas font and gray background
  • Markdown tables converted to properly formatted Word tables
  • Tab-indented bullet lists with proper Word formatting
  • Nested bullet lists (supports both tab and space indentation)
  • Horizontal rules (---) for section separation
  • Bold text (**bold**) and inline code (`code`) formatting

SUPPORTED MARKDOWN ELEMENTS:
----------------------------
- **Headers**: All levels (# through #####) with Amazon orange styling
- **Code Blocks**: Triple backtick syntax with language specification
- **Tables**: Markdown table syntax (| Column 1 | Column 2 |)
- **Bullet Lists**: Both - and * markers with tab indentation support
- **Numbered Lists**: Standard 1. 2. 3. format
- **Bold Text**: **bold text** formatting
- **Inline Code**: `inline code` with Consolas font
- **Horizontal Rules**: --- for section separation
- **Regular Paragraphs**: Standard text with Amazon Ember font

PROFESSIONAL DOCUMENT REQUIREMENTS:
-----------------------------------
For automatic professional conversion, markdown must contain:
- Document title using # format
- Document subtitle using ## format  
- Executive Summary using # Executive Summary format
- Proper AWS document structure (Notices, Abstract, Introduction sections)

REQUIREMENTS:
-------------
- Python 3.6 or higher
- python-docx library: pip install python-docx
- requests library: pip install requests (for AWS logo download)
- Read permissions on input Markdown file
- Write permissions in output directory
- Internet connection (for initial AWS logo download)

PERMISSIONS NEEDED:
------------------
- File system read access to input .md file
- File system write access to output directory
- Network access to d0.awsstatic.com (for AWS logo download)
- No special system or administrative permissions required

USAGE:
------
python markdown_to_docx_converter.py <input_markdown_file> [output_docx_file]

Arguments:
  input_markdown_file   Path to the input .md file (required)
  output_docx_file      Path for output .docx file (optional)
                        If not provided, uses input filename with .docx extension

EXAMPLES:
---------

Basic conversion (auto-generates output filename):
    python markdown_to_docx_converter.py document.md
    # Creates: document.docx with professional or generic formatting

Professional AWS document conversion:
    python markdown_to_docx_converter.py aws_migration_strategy.md
    # Creates: aws_migration_strategy.docx with 6-page front matter

Specify custom output filename:
    python markdown_to_docx_converter.py report.md final_report.docx

Convert with full paths:
    python markdown_to_docx_converter.py /path/to/input.md /path/to/output.docx

Windows usage:
    python markdown_to_docx_converter.py "C:\\Documents\\report.md" "C:\\Output\\report.docx"

INPUT FILE FORMAT:
------------------
Standard Markdown (.md) files with UTF-8 encoding.
Supports AWS Professional Services document structure and standard markdown syntax.

OUTPUT FORMAT:
--------------
**Professional AWS Documents** (when Executive Summary detected):
- 6-page front matter with title, notices, TOC, abstract, introduction
- Official AWS "Powered by AWS" logo on title page and footers
- Amazon Ember font throughout document
- Amazon orange headings (RGB: 255, 153, 0)
- Professional page headers starting on page 8
- Smart page breaks and section management
- Complete AWS Professional Services formatting

**Generic Documents** (standard markdown):
- Clean markdown to Word conversion
- Amazon Ember font and orange headings
- Standard page breaks for major sections
- No professional front matter

AWS LOGO INTEGRATION:
--------------------
- **Official Source**: Downloads from d0.awsstatic.com (AWS Co-Marketing Tools)
- **Title Page**: Vertically centered logo (2.5" width)
- **Page Footers**: Right-aligned logo (1.0" width) for content sections
- **Compliance**: AWS Trademark Guidelines compliant
- **Caching**: Local caching to avoid repeated downloads
- **Error Handling**: Graceful fallback if logo download fails

ERROR HANDLING:
---------------
- Missing input file: Clear error message with file path
- Invalid Markdown syntax: Graceful handling, continues processing
- Write permission issues: Error message with suggested solutions
- Large files: Processes efficiently without memory issues
- AWS logo download failures: Continues conversion without logo
- Network connectivity issues: Graceful degradation

LIMITATIONS:
------------
- Images in Markdown are not converted (text only)
- Complex HTML within Markdown is not processed
- Advanced Markdown extensions not supported
- Font availability depends on system (falls back to defaults)
- AWS logo requires internet connection for initial download

INTEGRATION:
------------
- **Template Generator Integration**: Called automatically by create_aws_markdown_template.py
- **Standalone Usage**: Can be used independently for any markdown file
- **AWS Documentation Workflow**: Part of complete AWS documentation system
- **Quality Assurance**: Built-in validation for professional document structure

TROUBLESHOOTING:
----------------
If conversion fails:
1. Check input file exists and is readable
2. Verify output directory has write permissions
3. Ensure python-docx is installed: pip install python-docx
4. Ensure requests library is installed: pip install requests
5. Check file encoding is UTF-8
6. Verify internet connection for AWS logo download
7. Check AWS document structure for professional conversion

AUTHOR: AWS Professional Services
VERSION: 3.0 - Professional AWS Document Converter
DATE: 2025
"""

import re
import requests
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn

def download_aws_logo():
    """
    Download official AWS 'Powered by AWS' logo from AWS Co-Marketing Tools page
    
    Official source: https://aws.amazon.com/co-marketing/
    Logo URL from AWS's official static assets domain: d0.awsstatic.com
    This is the officially approved logo for AWS customers per AWS Trademark Guidelines
    """
    # Official AWS 'Powered by AWS' logo from AWS Co-Marketing Tools page
    logo_url = "https://d0.awsstatic.com/logos/powered-by-aws.png"
    logo_filename = "aws_logo.png"
    
    # Check if logo already exists locally
    if os.path.exists(logo_filename):
        print(f"Official AWS logo already exists locally: {logo_filename}")
        return logo_filename
    
    try:
        print("Downloading official AWS 'Powered by AWS' logo from d0.awsstatic.com...")
        print("Source: AWS Co-Marketing Tools (https://aws.amazon.com/co-marketing/)")
        response = requests.get(logo_url, timeout=10)
        response.raise_for_status()
        
        with open(logo_filename, 'wb') as f:
            f.write(response.content)
        
        print(f"✅ Official AWS logo downloaded successfully: {logo_filename}")
        print("   This is the officially approved 'Powered by AWS' logo per AWS Trademark Guidelines")
        return logo_filename
    
    except requests.exceptions.RequestException as e:
        print(f"⚠️  Failed to download official AWS logo: {e}")
        print("   Continuing without logo...")
        return None
    except Exception as e:
        print(f"⚠️  Error saving official AWS logo: {e}")
        print("   Continuing without logo...")
        return None

def add_logo_to_title_page(doc, logo_path):
    """Add AWS logo to the center of the title page (vertically and horizontally centered)"""
    if not logo_path or not os.path.exists(logo_path):
        return
    
    try:
        # Find the position after subtitle to insert the logo
        paragraphs = list(doc.paragraphs)
        
        # Look for title and subtitle paragraphs to insert logo after them
        title_found = False
        subtitle_found = False
        insert_position = 2  # Default position after title and subtitle
        
        for i, paragraph in enumerate(paragraphs):
            if paragraph.text and len(paragraph.text.strip()) > 0:
                if not title_found:
                    title_found = True
                elif not subtitle_found:
                    subtitle_found = True
                    insert_position = i + 1
                    break
        
        # Insert logo in the middle of the page (after title/subtitle, before date)
        # Calculate middle position - we want the logo roughly in the center
        middle_position = len(paragraphs) // 2
        
        # Create logo paragraph at the middle position
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(2.5))
        
        # Move the logo paragraph to the middle position
        logo_element = logo_paragraph._element
        parent = logo_element.getparent()
        parent.remove(logo_element)
        parent.insert(middle_position, logo_element)
        
        print("✅ AWS logo added to title page (vertically centered)")
    except Exception as e:
        print(f"⚠️  Failed to add logo to title page: {e}")

def add_logo_to_footer(section, logo_path):
    """Add AWS logo to the left side and page numbers to the right side of the footer"""
    if not logo_path or not os.path.exists(logo_path):
        return
    
    try:
        # Access the footer
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Clear existing footer content
        for paragraph in footer.paragraphs:
            paragraph.clear()
        
        # Create footer paragraph with both logo and page numbers
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Set up tab stops BEFORE adding content
        tab_stops = footer_para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        
        # Add AWS logo (left-aligned - stays at left margin)
        logo_run = footer_para.add_run()
        logo_run.add_picture(logo_path, width=Inches(1.0))  # Smaller size for footer
        
        # Add tab character to move to right side
        tab_run = footer_para.add_run('\t')
        
        # Add page numbers in "1 of X" format (will appear at right tab stop)
        page_run = footer_para.add_run()
        page_run.font.name = 'Amazon Ember'
        page_run.font.size = Pt(10)
        
        # Add page number field (Word will auto-calculate)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        page_run._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE \\* MERGEFORMAT'
        page_run._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_run._element.append(fldChar2)
        
        # Add " of " text
        of_run = footer_para.add_run(' of ')
        of_run.font.name = 'Amazon Ember'
        of_run.font.size = Pt(10)
        
        # Add total pages field
        total_run = footer_para.add_run()
        total_run.font.name = 'Amazon Ember'
        total_run.font.size = Pt(10)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        total_run._element.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = 'NUMPAGES \\* MERGEFORMAT'
        total_run._element.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        total_run._element.append(fldChar4)
        
        print("✅ AWS logo (left-aligned) and page numbers (1 of X format, right-aligned) added to footer")
    except Exception as e:
        print(f"⚠️  Failed to add logo and page numbers to footer: {e}")

def customize_built_in_styles(doc):
    """Customize built-in Word styles with Amazon Ember font"""
    
    # Customize built-in heading styles
    for i in range(1, 6):
        heading_style = doc.styles[f'Heading {i}']
        # Set font properties on the character format
        heading_style.font.name = 'Amazon Ember'
        heading_style.font.bold = True
        # Remove space before for Heading 3 (Executive Summary)
        if i == 3:
            heading_style.paragraph_format.space_before = Pt(0)
            heading_style.paragraph_format.space_after = Pt(6)
        # Also ensure the base character format is set
        if hasattr(heading_style, '_element'):
            rpr = heading_style._element.get_or_add_rPr()
            # Set font family
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts')
                rpr.append(rfonts)
            rfonts.set(qn('w:ascii'), 'Amazon Ember')
            rfonts.set(qn('w:hAnsi'), 'Amazon Ember')
            rfonts.set(qn('w:cs'), 'Amazon Ember')
    
    # Customize Normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Amazon Ember'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    # Also ensure the base character format is set for Normal
    if hasattr(normal_style, '_element'):
        rpr = normal_style._element.get_or_add_rPr()
        # Set font family
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), 'Amazon Ember')
        rfonts.set(qn('w:hAnsi'), 'Amazon Ember')
        rfonts.set(qn('w:cs'), 'Amazon Ember')
    
    # Code style
    code_style = doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text):
    """Add a code block with proper formatting"""
    # Add code block with background shading
    p = doc.add_paragraph(style='CustomCode')
    run = p.runs[0] if p.runs else p.add_run()
    run.text = code_text
    
    # Add light gray shading to simulate code block background
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(shading_elm)

def add_table(doc, table_data):
    """Add a properly formatted table to the document"""
    if not table_data or len(table_data) < 2:
        return
    
    # Create table with proper dimensions
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    # Set table formatting
    table.autofit = False
    table.allow_autofit = True
    
    # Fill table data
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_data.strip()
            
            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Amazon Ember'
                        run.font.bold = True
                        run.font.size = Pt(11)
            else:
                # Format data rows
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Amazon Ember'
                        run.font.size = Pt(10)
    
    # Add spacing after table
    doc.add_paragraph()

def process_markdown_line(doc, line, in_code_block, code_buffer, is_professional_doc=False, table_buffer=None):
    """Process a single line of markdown and add to document"""
    
    if table_buffer is None:
        table_buffer = []
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block
            if code_buffer:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer.clear()
            return False, code_buffer, table_buffer
        else:
            # Start of code block
            return True, code_buffer, table_buffer
    
    if in_code_block:
        code_buffer.append(line)
        return True, code_buffer, table_buffer
    
    # Handle markdown tables
    if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
        # Check if this is a table separator line (contains dashes)
        if '---' in line or '-' * 3 in line:
            # This is a table separator line - skip it completely
            return False, code_buffer, table_buffer
        else:
            # This is a table row
            cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]  # Remove empty first/last elements
            table_buffer.append(cells)
            return False, code_buffer, table_buffer
    elif table_buffer and not ('|' in line and line.strip().startswith('|')):
        # End of table - process the buffered table
        if len(table_buffer) > 0:
            add_table(doc, table_buffer)
            table_buffer.clear()
        # Continue processing this line normally (fall through)
    
    # Skip horizontal rules
    if line.strip() == '---':
        doc.add_paragraph()  # Add spacing
        return False, code_buffer, table_buffer
    
    # Handle headings
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        
        # For professional document, disable automatic page breaks since we manage sections manually
        # Only add page breaks for non-professional documents
        if not is_professional_doc and level <= 2:
            doc.add_page_break()
        
        if level == 1:
            p = doc.add_paragraph(text, style='Heading 1')
        elif level == 2:
            p = doc.add_paragraph(text, style='Heading 2')
        elif level == 3:
            p = doc.add_paragraph(text, style='Heading 3')
        elif level == 4:
            p = doc.add_paragraph(text, style='Heading 4')
        elif level == 5:
            p = doc.add_paragraph(text, style='Heading 5')
        else:
            p = doc.add_paragraph(text, style='Heading 5')
        
        # Apply Amazon orange color to content headings only
        for run in p.runs:
            run.font.name = 'Amazon Ember'
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange for content headings
        
        return False, code_buffer, table_buffer
    
    # Handle bullet points (including tab-indented ones)
    # Only process lines that actually start with bullet markers (- or * followed by space)
    stripped_line = line.lstrip()
    if (stripped_line.startswith('- ') or stripped_line.startswith('* ') or 
        stripped_line == '-' or stripped_line == '*'):
        # Extract the bullet text (everything after the bullet marker)
        stripped_line = line.lstrip()
        text = stripped_line[1:].strip()
        
        # Calculate indent level based on leading whitespace before the bullet marker
        leading_whitespace = line[:len(line) - len(line.lstrip())]
        tab_count = leading_whitespace.count('\t')
        space_count = leading_whitespace.count(' ')
        
        # Determine indent level based on actual indentation in the markdown:
        # - Lines with no indentation before bullet = level 0 (top-level)
        # - Lines with tab indentation before bullet = level based on tab count
        # - Lines with space indentation before bullet = level based on space count (2 spaces per level)
        if len(leading_whitespace) == 0:
            indent_level = 0  # Top-level bullet (no indentation)
        elif tab_count > 0:
            indent_level = tab_count  # Tab-based indentation (each tab = 1 level)
        elif space_count > 0:
            indent_level = space_count // 2  # Space-based indentation (2 spaces per level)
        else:
            indent_level = 0  # Default to top-level
        
        # Create proper Word bullet list with appropriate indentation
        if indent_level == 0:
            p = doc.add_paragraph(style='List Bullet')
        else:
            # For nested bullets, use Normal style and manually add bullet and indentation
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5 * indent_level)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            # Add bullet character manually
            bullet_run = p.add_run('• ')
            bullet_run.font.name = 'Amazon Ember'
        
        # Process bold text in bullets
        if '**' in text:
            parts = text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Amazon Ember'
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True
        else:
            run = p.add_run(text)
            run.font.name = 'Amazon Ember'
        return False, code_buffer, table_buffer
    
    # Handle script titles (e.g., **`script.sh`**) as Heading 4 - MUST come before numbered lists
    script_title_pattern = r'^\*\*`([^`]+\.sh)`\*\*$'
    script_match = re.match(script_title_pattern, line.strip())
    if script_match:
        script_name = script_match.group(1)
        p = doc.add_paragraph(script_name, style='Heading 4')
        # Apply Amazon orange color and Amazon Ember font to script titles
        for run in p.runs:
            run.font.name = 'Amazon Ember'
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
        return False, code_buffer, table_buffer
    
    # Handle numbered lists
    if re.match(r'^\d+\.', line.strip()):
        text = re.sub(r'^\d+\.\s*', '', line.strip())
        # Create proper Word numbered list
        p = doc.add_paragraph(style='List Number')
        
        # Process bold text in numbered lists
        if '**' in text:
            parts = text.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Amazon Ember'
                if i % 2 == 1:  # Odd indices are bold
                    run.bold = True
        else:
            run = p.add_run(text)
            run.font.name = 'Amazon Ember'
        return False, code_buffer, table_buffer
    
    # Handle regular paragraphs
    if line.strip():
        p = doc.add_paragraph(style='Normal')
        
        # Process bold text and inline code
        text = line.strip()
        
        # Handle inline code first (backticks)
        code_pattern = r'`([^`]+)`'
        parts = re.split(code_pattern, text)
        
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Odd indices are code
                run = p.add_run(part)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
            else:
                # Handle bold text in non-code parts
                if '**' in part:
                    bold_parts = part.split('**')
                    for j, bold_part in enumerate(bold_parts):
                        run = p.add_run(bold_part)
                        if j % 2 == 1:  # Odd indices are bold
                            run.bold = True
                else:
                    p.add_run(part)
    else:
        # Empty line - add paragraph break
        doc.add_paragraph()
    
    return False, code_buffer, table_buffer

def create_title_page(doc, title_text, subtitle_text):
    """Create professional title page - single page"""
    # Title starts at the top (first line)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_p.add_run(title_text)
    title_run.font.name = 'Amazon Ember'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
    
    # Subtitle
    if subtitle_text:
        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle_p.add_run(subtitle_text)
        subtitle_run.font.name = 'Amazon Ember'
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.bold = True
    
    # Add spacing to push date to bottom
    for _ in range(15):
        doc.add_paragraph()
    
    # Date - bottom of page
    from datetime import datetime
    current_date = datetime.now().strftime("%B %Y")
    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(current_date)
    date_run.font.name = 'Amazon Ember'
    date_run.font.size = Pt(14)
    
    # Last line - page break
    doc.add_page_break()

def create_notices_page(doc):
    """Create notices and disclaimers page"""
    # Notices heading with Amazon orange color
    notices_p = doc.add_paragraph("Notices", style='Heading 1')
    # Apply Amazon orange color to Notices heading
    for run in notices_p.runs:
        run.font.name = 'Amazon Ember'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
    
    # AWS disclaimer with current year
    from datetime import datetime
    current_year = datetime.now().year
    disclaimer_text = f"""This document is provided for informational purposes only. It represents AWS's current product offerings and practices as of the date of issue of this document, which are subject to change without notice. Customers are responsible for making their own independent assessment of the information in this document and any use of AWS's products or services, each of which is provided "as is" without warranty of any kind, whether express or implied. This document does not create any warranties, representations, contractual commitments, conditions or assurances from AWS, its affiliates, suppliers or licensors. The responsibilities and liabilities of AWS to its customers are controlled by AWS agreements, and this document is not part of, nor does it modify, any agreement between AWS and its customers.

© {current_year} Amazon Web Services, Inc. or its affiliates. All rights reserved."""
    
    disclaimer_p = doc.add_paragraph(disclaimer_text, style='Normal')
    
    # Add page break
    doc.add_page_break()

def create_table_of_contents_pages(doc):
    """Create table of contents pages (2 pages reserved)"""
    # TOC heading with Amazon orange color
    toc_p = doc.add_paragraph("Table of Contents", style='Heading 1')
    # Apply Amazon orange color to Table of Contents heading
    for run in toc_p.runs:
        run.font.name = 'Amazon Ember'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
    
    # Placeholder text
    placeholder_p = doc.add_paragraph("Table of Contents will be generated automatically by Microsoft Word.", style='Normal')
    placeholder_p2 = doc.add_paragraph("To update: Right-click → Update Field → Update entire table", style='Normal')
    
    # Add page break for first TOC page
    doc.add_page_break()
    
    # Second TOC page (reserved space)
    doc.add_paragraph("(Table of Contents continued)", style='Normal')
    
    # Add page break
    doc.add_page_break()

def create_abstract_page(doc, abstract_text=None):
    """Create abstract page"""
    # Abstract heading with Amazon orange color
    abstract_p = doc.add_paragraph("Abstract", style='Heading 1')
    # Apply Amazon orange color to Abstract heading
    for run in abstract_p.runs:
        run.font.name = 'Amazon Ember'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
    
    # Abstract content - use provided text or placeholder
    if abstract_text:
        abstract_content_p = doc.add_paragraph(abstract_text, style='Normal')
    else:
        placeholder_text = "This document provides comprehensive information and recommendations. The abstract will be automatically generated based on the document content or can be manually updated."
        abstract_content_p = doc.add_paragraph(placeholder_text, style='Normal')
    
    # Add page break
    doc.add_page_break()

def create_introduction_page(doc, introduction_text=None):
    """Create introduction page"""
    # Introduction heading with Amazon orange color
    intro_p = doc.add_paragraph("Introduction", style='Heading 1')
    # Apply Amazon orange color to Introduction heading
    for run in intro_p.runs:
        run.font.name = 'Amazon Ember'
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
    
    # Introduction content - use provided text or placeholder
    if introduction_text:
        intro_content_p = doc.add_paragraph(introduction_text, style='Normal')
    else:
        placeholder_text = "This document provides detailed information and guidance. The introduction section will be automatically generated based on the document content or can be manually updated to reflect the specific context and objectives."
        intro_content_p = doc.add_paragraph(placeholder_text, style='Normal')
    
    # No page break here - Executive Summary will handle its own page positioning

def clear_section_header(section):
    """Clear header from a section"""
    section.header.is_linked_to_previous = False
    # Clear existing header content
    for paragraph in section.header.paragraphs:
        paragraph.clear()

def add_header_to_section(section, header_text):
    """Add header to a specific section"""
    section.header.is_linked_to_previous = False
    header_para = section.header.paragraphs[0]
    header_para.text = header_text
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Style the header - use black for page headers
    for run in header_para.runs:
        run.font.name = 'Amazon Ember'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black for page headers

def extract_section_content(lines, section_heading):
    """Extract content from a specific section (e.g., '## Abstract')"""
    content_lines = []
    in_section = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Check if we're starting the target section
        if line_stripped == section_heading:
            in_section = True
            continue
        
        # Check if we've hit another section heading (end of current section)
        if in_section and (line_stripped.startswith('# ') or line_stripped.startswith('## ') or line_stripped.startswith('### ')):
            break
        
        # Collect content lines (skip horizontal rules and empty lines at start)
        if in_section:
            if line_stripped == '---':
                continue
            if line_stripped or content_lines:  # Start collecting after first non-empty line
                content_lines.append(line.rstrip('\n'))
    
    # Join lines and clean up
    content = '\n'.join(content_lines).strip()
    return content if content else None

def convert_markdown_to_docx(md_file_path, output_path):
    """Convert markdown file to DOCX format with professional structure"""
    
    # Download AWS logo for professional documents
    logo_path = download_aws_logo()
    
    # Create new document
    doc = Document()
    
    # Set up page margins for all sections
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Customize built-in styles
    customize_built_in_styles(doc)
    
    # Read markdown file first to extract title information
    try:
        with open(md_file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
    except FileNotFoundError:
        print(f"Error: File '{md_file_path}' not found.")
        return False
    except Exception as e:
        print(f"Error reading file: {e}")
        return False
    
    # Extract title and subtitle from markdown
    title_text = None
    subtitle_text = None
    
    for line in lines:
        line = line.strip()
        if line.startswith('# ') and not title_text:
            title_text = line[2:].strip()
        elif line.startswith('## ') and not subtitle_text:
            subtitle_text = line[3:].strip()
        elif title_text and subtitle_text:
            break
    
    # Extract Abstract and Introduction content from markdown
    abstract_content = extract_section_content(lines, '## Abstract')
    introduction_content = extract_section_content(lines, '## Introduction')
    
    # Determine if this should be a professional document based on structure
    # Look for professional document indicators
    has_executive_summary = any('# Executive Summary' in line for line in lines)
    is_professional_doc = has_executive_summary and title_text and subtitle_text
    
    if is_professional_doc:
        # Create professional document structure
        print("Detected professional document structure - applying full professional format...")
        print("Creating title page...")
        create_title_page(doc, title_text, subtitle_text)
        
        # Add AWS logo to title page
        if logo_path:
            add_logo_to_title_page(doc, logo_path)
        
        print("Creating notices page...")
        create_notices_page(doc)
        
        print("Creating table of contents pages...")
        create_table_of_contents_pages(doc)
        
        print("Creating abstract page...")
        create_abstract_page(doc, abstract_content)
        
        print("Creating introduction page...")
        create_introduction_page(doc, introduction_content)
        
        # Clear headers for pages 1-6 (title, notices, TOC, abstract, introduction)
        for section in doc.sections:
            clear_section_header(section)
    else:
        print("Processing generic markdown document...")
    
    # Read markdown file
    try:
        with open(md_file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
    except FileNotFoundError:
        print(f"Error: File '{md_file_path}' not found.")
        return False
    except Exception as e:
        print(f"Error reading file: {e}")
        return False
    
    print("Processing markdown content...")
    
    # Process markdown content
    in_code_block = False
    code_buffer = []
    table_buffer = []
    executive_summary_processed = False
    content_section_started = False
    title_lines_skipped = 0  # Track how many title lines we've skipped
    skipping_section = None  # Track which section we're currently skipping
    
    for line_num, line in enumerate(lines):
        line = line.rstrip('\n')
        
        # For professional document, skip the first two heading lines (title and subtitle)
        # since they're already handled by the title page
        if is_professional_doc and not content_section_started:
            if line.strip() == f"# {title_text}" and title_lines_skipped == 0:
                title_lines_skipped += 1
                continue
            elif line.strip() == f"## {subtitle_text}" and title_lines_skipped == 1:
                title_lines_skipped += 1
                continue
        
        # Skip the Notices, Abstract, and Introduction sections since they're already processed in the front matter
        if is_professional_doc and not content_section_started:
            if line.strip() == "## Notices":
                skipping_section = "Notices"
                print(f"Skipping {skipping_section} section - already processed in front matter")
                continue
            elif line.strip() == "## Abstract":
                skipping_section = "Abstract"
                print(f"Skipping {skipping_section} section - already processed in front matter")
                continue
            elif line.strip() == "## Introduction":
                skipping_section = "Introduction"
                print(f"Skipping {skipping_section} section - already processed in front matter")
                continue
            elif skipping_section and (line.strip().startswith('## ') or line.strip().startswith('### ') or line.strip().startswith('# ')):
                # We've hit the next section, stop skipping
                print(f"Finished skipping {skipping_section} section")
                skipping_section = None
                # Don't continue here, process this line normally
            elif skipping_section:
                # We're still in a section we're skipping
                continue
        
        # Handle Executive Summary specially for professional document - this should start the content section WITHOUT headers yet
        if line.strip() == "# Executive Summary" and not executive_summary_processed and is_professional_doc:
            print("Processing Executive Summary - starting on its own page (page 7)...")
            
            # Add page break to ensure Executive Summary starts on its own page (page 7)
            doc.add_page_break()
            
            # Add Executive Summary heading
            exec_summary_p = doc.add_paragraph("Executive Summary", style='Heading 1')
            # Ensure proper spacing for the heading
            exec_summary_p.paragraph_format.space_before = Pt(0)
            exec_summary_p.paragraph_format.space_after = Pt(6)
            # Apply Amazon orange color to Executive Summary heading
            for run in exec_summary_p.runs:
                run.font.name = 'Amazon Ember'
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange
            
            executive_summary_processed = True
            content_section_started = True
            
            # Continue processing normally - don't call process_markdown_line for this heading since we handled it
            continue
        # Process Executive Summary content lines (including ## subsections within Executive Summary)
        elif content_section_started == True and is_professional_doc:
            # Check if this is the end of Executive Summary (next major section that's not a subsection)
            # Executive Summary ends when we hit a major section that's not part of the Executive Summary
            if (line.startswith('## ') and 
                line.strip() not in ['## Business Value and Strategic Importance', 
                                   '## Operational Benefits and Risk Mitigation', 
                                   '## Implementation Approach and Success Factors']):
                print(f"Executive Summary content complete - adding headers for main content sections")
                print(f"Adding headers starting from first major section: {line.strip()}")
                
                # Create ONE new section with headers for all subsequent content
                new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.top_margin = Inches(1)
                new_section.bottom_margin = Inches(1)
                new_section.left_margin = Inches(1)
                new_section.right_margin = Inches(1)
                
                # Add header using the document title
                header_text = f"AWS Professional Services - {title_text}"
                add_header_to_section(new_section, header_text)
                
                # Add AWS logo to footer for content sections
                if logo_path:
                    add_logo_to_footer(new_section, logo_path)
                
                # Manually process the heading to ensure it appears at the top of the new section
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                
                # Add the heading directly to the document (which will now be in the new section)
                if level == 2:
                    p = doc.add_paragraph(text, style='Heading 2')
                else:
                    p = doc.add_paragraph(text, style='Heading 2')  # Force level 2 for ## headings
                
                # Apply Amazon orange color to the heading
                for run in p.runs:
                    run.font.name = 'Amazon Ember'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 153, 0)  # Amazon orange for content headings
                
                # Mark that we've added headers so subsequent ## headings use regular page breaks
                content_section_started = "headers_added"
            else:
                # Process Executive Summary content normally (including ## subsections)
                in_code_block, code_buffer, table_buffer = process_markdown_line(doc, line, in_code_block, code_buffer, is_professional_doc, table_buffer)
        # For subsequent major sections, add smart page breaks
        elif content_section_started == "headers_added" and line.startswith('## ') and is_professional_doc:
            print(f"Processing subsequent major section with smart page break: {line.strip()}")
            
            # Add a paragraph break to create some separation, but let Word handle natural page flow
            # This provides visual separation without forcing unnecessary page breaks
            doc.add_paragraph()
            
            # Process the heading normally
            in_code_block, code_buffer, table_buffer = process_markdown_line(doc, line, in_code_block, code_buffer, is_professional_doc, table_buffer)
        else:
            # Process the line normally
            in_code_block, code_buffer, table_buffer = process_markdown_line(doc, line, in_code_block, code_buffer, is_professional_doc, table_buffer)
    
    # Handle any remaining code block
    if code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))
    
    # Handle any remaining table buffer
    if table_buffer:
        add_table(doc, table_buffer)
    
    # Save document
    try:
        doc.save(output_path)
        print(f"Successfully converted '{md_file_path}' to '{output_path}'")
        return True
    except Exception as e:
        print(f"Error saving document: {e}")
        return False

def main():
    """Main function"""
    import sys
    import os
    
    # Check for command line arguments
    if len(sys.argv) < 2:
        print("Usage: python markdown_to_docx_converter.py <input_markdown_file> [output_docx_file]")
        print("\nExample:")
        print("  python markdown_to_docx_converter.py document.md")
        print("  python markdown_to_docx_converter.py document.md output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"❌ Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    # Generate output filename if not provided
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        # Generate output filename by replacing .md extension with .docx
        base_name = os.path.splitext(input_file)[0]
        output_file = f"{base_name}.docx"
    
    print("Converting Markdown to DOCX...")
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    print("-" * 50)
    
    success = convert_markdown_to_docx(input_file, output_file)
    
    if success:
        print("-" * 50)
        print("✅ Conversion completed successfully!")
        print(f"📄 Word document created: {output_file}")
        
        # Determine document type by checking for professional structure indicators
        try:
            with open(input_file, 'r', encoding='utf-8') as file:
                content = file.read()
                has_executive_summary = '# Executive Summary' in content
                has_title_structure = content.count('# ') >= 1 and content.count('## ') >= 1
                is_professional = has_executive_summary and has_title_structure
        except:
            is_professional = False
        
        if is_professional:
            print("\n🎯 Professional Document Structure:")
            print("• Title Page - Extracted from markdown title and subtitle")
            print("• Notices Page - AWS disclaimers and copyright")
            print("• Table of Contents Pages (2 pages reserved)")
            print("• Abstract Page - Uses content from markdown ## Abstract section")
            print("• Introduction Page - Uses content from markdown ## Introduction section")
            print("• Content Section - Starts with Executive Summary")
            print("\n📋 Professional Features:")
            print("• Official AWS 'Powered by AWS' logo on title page and footers")
            print("• Black page headers starting on page 8")
            print("• Intelligent page breaks for major sections")
            print("• Amazon orange content headings")
            print("• Complete professional document structure")
        else:
            print("\n📝 Generic Document Features:")
            print("• Clean markdown conversion")
            print("• No professional front matter")
            print("• Standard formatting only")
        
        print("\n🎨 Universal Formatting Features:")
        print("• Amazon Ember font throughout document")
        print("• Tab-indented bullet point support")
        print("• Code blocks with Consolas font and gray background")
        print("• Microsoft Word built-in styles (Heading 1-5, Normal)")
        print("• Bold text and inline code formatting")
    else:
        print("❌ Conversion failed!")

if __name__ == "__main__":
    main()
